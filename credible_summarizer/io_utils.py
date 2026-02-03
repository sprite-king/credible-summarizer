import csv
import html
import json
from datetime import datetime, timezone
from io import BytesIO, StringIO
from typing import Dict, List

try:
    from docx import Document
except Exception:  # pragma: no cover - optional dependency
    Document = None

try:
    from PyPDF2 import PdfReader
except Exception:  # pragma: no cover - optional dependency
    PdfReader = None

from .models import Citation
from .text_utils import format_score


def extract_text_from_upload(uploaded_file) -> str:
    if uploaded_file is None:
        return ""
    file_name = uploaded_file.name.lower()
    if file_name.endswith(".txt"):
        return uploaded_file.read().decode("utf-8", errors="ignore")
    if file_name.endswith(".pdf"):
        if PdfReader is None:
            return ""
        reader = PdfReader(BytesIO(uploaded_file.read()))
        pages = [page.extract_text() or "" for page in reader.pages]
        return "\n".join(pages)
    if file_name.endswith(".docx"):
        if Document is None:
            return ""
        doc = Document(BytesIO(uploaded_file.read()))
        return "\n".join(p.text for p in doc.paragraphs)
    if file_name.endswith(".doc"):
        return ""
    return ""


def export_markdown(
    sentences: List[str],
    summary_indices: List[int],
    citations: List[List[Citation]],
    uncertainty_flags: List[bool],
) -> str:
    lines = ["# 可信摘要（带引用）", ""]
    for i, idx in enumerate(summary_indices):
        flag = "⚠️" if uncertainty_flags[i] else "✅"
        lines.append(f"- {flag} {sentences[idx]}")
        for cite in citations[i]:
            lines.append(f"  - 引用（{format_score(cite.score)}）: {cite.sentence}")
        lines.append("")
    return "\n".join(lines).strip()


def export_html(
    sentences: List[str],
    summary_indices: List[int],
    citations: List[List[Citation]],
    uncertainty_flags: List[bool],
) -> str:
    lines = ["<h1>可信摘要（带引用）</h1>", "<ul>"]
    for i, idx in enumerate(summary_indices):
        flag = "⚠️" if uncertainty_flags[i] else "✅"
        lines.append(f"<li>{flag} {html.escape(sentences[idx])}")
        lines.append("<ul>")
        for cite in citations[i]:
            lines.append(
                f"<li>引用（{format_score(cite.score)}）: {html.escape(cite.sentence)}</li>"
            )
        lines.append("</ul></li>")
    lines.append("</ul>")
    return "\n".join(lines)


def export_docx(
    sentences: List[str],
    summary_indices: List[int],
    citations: List[List[Citation]],
    uncertainty_flags: List[bool],
) -> bytes:
    if Document is None:
        return b""
    doc = Document()
    doc.add_heading("可信摘要（带引用）", level=1)
    for i, idx in enumerate(summary_indices):
        flag = "⚠️" if uncertainty_flags[i] else "✅"
        doc.add_paragraph(f"{flag} {sentences[idx]}", style="List Bullet")
        for cite in citations[i]:
            doc.add_paragraph(
                f"引用（{format_score(cite.score)}）: {cite.sentence}", style="List Bullet 2"
            )
    output = BytesIO()
    doc.save(output)
    return output.getvalue()


def export_json(
    sentences: List[str],
    summary_indices: List[int],
    citations: List[List[Citation]],
    uncertainty_flags: List[bool],
) -> str:
    items = []
    for i, idx in enumerate(summary_indices):
        items.append(
            {
                "summary": sentences[idx],
                "uncertain": bool(uncertainty_flags[i]),
                "citations": [
                    {"sentence": c.sentence, "score": c.score, "index": c.index}
                    for c in citations[i]
                ],
            }
        )
    return json.dumps({"items": items}, ensure_ascii=False, indent=2)


def export_csv(
    sentences: List[str],
    summary_indices: List[int],
    citations: List[List[Citation]],
    uncertainty_flags: List[bool],
) -> str:
    output = StringIO()
    writer = csv.writer(output)
    writer.writerow(["summary", "uncertain", "citation_sentence", "citation_score", "citation_index"])
    for i, idx in enumerate(summary_indices):
        for c in citations[i]:
            writer.writerow([sentences[idx], int(uncertainty_flags[i]), c.sentence, format_score(c.score), c.index])
    return output.getvalue()


def export_experiment_record(params: Dict, versions: Dict) -> str:
    record = {
        "timestamp": datetime.now(timezone.utc).isoformat(),
        "params": params,
        "versions": versions,
    }
    return json.dumps(record, ensure_ascii=False, indent=2)
